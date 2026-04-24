#!/usr/bin/env python3
"""
Script to create a Word document with all content from the MoonShot app screens
"""

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("python-docx is not installed. Installing...")
    import subprocess
    import sys
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

def add_heading(doc, text, level=1):
    """Add a heading with consistent styling"""
    heading = doc.add_heading(text, level=level)
    heading.style.font.name = 'Inter'
    heading.style.font.size = Pt(18 if level == 1 else 16 if level == 2 else 14)
    return heading

def add_paragraph(doc, text, bold=False, italic=False, color=None):
    """Add a paragraph with consistent styling"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Inter'
    run.font.size = Pt(11)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def add_bullet(doc, text, level=0):
    """Add a bullet point"""
    p = doc.add_paragraph(text, style='List Bullet' if level == 0 else 'List Bullet 2')
    p.style.font.name = 'Inter'
    p.style.font.size = Pt(11)
    return p

# Create document
doc = Document()

# Set default font
style = doc.styles['Normal']
style.font.name = 'Inter'
style.font.size = Pt(11)

# Title
title = doc.add_heading('Cognite MoonShot - Content Outline', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Home Page
add_heading(doc, 'Home Page', 1)
add_paragraph(doc, 'Cognite', bold=True)
add_paragraph(doc, 'MoonShot', bold=True)
add_paragraph(doc, '100 × 35')
add_paragraph(doc, '$100 Billion in realize customer value by 2035')
add_paragraph(doc, '')

add_paragraph(doc, 'Path Options:', bold=True)
add_bullet(doc, 'What is Moonshot? - Understand the mission')
add_bullet(doc, 'Why it matters - Discover the impact')
add_bullet(doc, 'How we get there - Explore the roadmap')
add_paragraph(doc, '')

# Path 1: What is Moonshot
add_heading(doc, 'Path 1: What is Moonshot?', 1)

# Screen 1
add_heading(doc, 'Screen 1: What is MoonShot?', 2)
add_paragraph(doc, "MoonShot is our pledge to deliver $100 Billion of realized customer value by 2035.")
add_paragraph(doc, '')
add_paragraph(doc, 'Not just a slogan - It\'s how we differentiate', bold=True)
add_paragraph(doc, '')

add_paragraph(doc, 'Value Assurance', bold=True)
add_paragraph(doc, 'We guarantee measurable business outcomes, not just technology implementation.')
add_paragraph(doc, '')

add_paragraph(doc, 'Beyond Pilots', bold=True)
add_paragraph(doc, 'Moving from proof-of-concept to production-scale impact that transforms operations.')
add_paragraph(doc, '')

add_paragraph(doc, 'P&L Impact', bold=True)
add_paragraph(doc, 'Directly measurable financial results that appear in customer profit and loss statements.')
add_paragraph(doc, '')

# Screen 2
add_heading(doc, 'Screen 2: Real Customer Value', 2)
add_paragraph(doc, 'Track, quantify, and prove value – at scale – across our entire customer base.')
add_paragraph(doc, '')
add_paragraph(doc, 'Core Principles', bold=True)
add_paragraph(doc, '')

add_paragraph(doc, 'Customer value, not ARR.', bold=True)
add_paragraph(doc, "Impact for our customers' P&L (revenue, cost, risk, sustainability), not just internal metrics or feature usage.")
add_paragraph(doc, '')

add_paragraph(doc, 'Measured, not aspired.', bold=True)
add_paragraph(doc, 'We work with customers to quantify outcomes, not just "feel" the improvement.')
add_paragraph(doc, '')

add_paragraph(doc, 'Realized, not estimated.', bold=True)
add_paragraph(doc, "Value that has actually materialized in the customer's operations, not just in a slide or business case.")
add_paragraph(doc, '')

# Screen 3
add_heading(doc, 'Screen 3: 5 levels to track value', 2)
add_paragraph(doc, 'The Journey to Proof', bold=True)
add_paragraph(doc, '')

add_paragraph(doc, '1. Value Estimated', bold=True)
add_paragraph(doc, "Cognite's initial estimate of what a use case could deliver, used for the business case.")
add_paragraph(doc, '')

add_paragraph(doc, '2. Customer Agreed', bold=True)
add_paragraph(doc, 'The customer has reviewed and agreed to the value estimate.')
add_paragraph(doc, '')

add_paragraph(doc, '3. Value Realized', bold=True)
add_paragraph(doc, 'Unilateral estimate of value already realized based on early data post-deployment.')
add_paragraph(doc, '')

add_paragraph(doc, '4. Value Measured', bold=True)
add_paragraph(doc, "We formally measure realized value together with the customer. Results are confidential but robust.")
add_paragraph(doc, '')

add_paragraph(doc, '5. Value Made Public', bold=True)
add_paragraph(doc, 'The customer agrees to share the measured impact publicly – our gold standard of proof.')
add_paragraph(doc, '')

# Path 2: Why it matters
add_heading(doc, 'Path 2: Why it matters', 1)

# Screen 1
add_heading(doc, 'Screen 1: Why 100 × 35 matters', 2)
add_paragraph(doc, 'As AI and data become mission-critical for every industry, the cost of failure is no longer acceptable. Companies need partners who can deliver proven value, not just technology.')
add_paragraph(doc, '')
add_paragraph(doc, 'The Stark Reality', bold=True)
add_paragraph(doc, '')

add_paragraph(doc, 'The True Cost', bold=True)
add_paragraph(doc, 'This failure rate represents more than wasted investment—it\'s lost opportunity, competitive disadvantage, and eroded trust in AI/data initiatives.')
add_paragraph(doc, '')

add_bullet(doc, 'Billions in sunk costs with no ROI')
add_bullet(doc, 'Years of delayed digital transformation')
add_bullet(doc, 'Team burnout and initiative fatigue')
add_bullet(doc, 'Loss of competitive advantage')
add_paragraph(doc, '')

# Screen 2
add_heading(doc, 'Screen 2: Value assurance is our differentiator', 2)
add_paragraph(doc, 'While the industry struggles with a 95% failure rate, Cognite\'s value assurance approach ensures every initiative delivers measurable, scalable business impact.')
add_paragraph(doc, '')
add_paragraph(doc, 'Pillars of Value Assurance', bold=True)
add_paragraph(doc, '')

add_paragraph(doc, 'Technology-First → Value-First', bold=True)
add_paragraph(doc, 'Every decision starts with business impact, not technical capability.')
add_paragraph(doc, '')

add_paragraph(doc, 'Pilot Success → Scaled Impact', bold=True)
add_paragraph(doc, 'Success is measured by enterprise-wide adoption, not proof of concept.')
add_paragraph(doc, '')

add_paragraph(doc, 'Vendor Promises → Customer Validation', bold=True)
add_paragraph(doc, 'Value is validated by the customer, not claimed by the vendor.')
add_paragraph(doc, '')

add_paragraph(doc, 'Opaque Metrics → Transparent KPIs', bold=True)
add_paragraph(doc, 'Clear, measurable, publicly shareable evidence of business impact.')
add_paragraph(doc, '')

# Screen 3
add_heading(doc, 'Screen 3: Culture Shift', 2)
add_paragraph(doc, '100 × 35 changes how we think and work internally. It says that "value" is not one team\'s responsibility – it\'s our shared operating system.')
add_paragraph(doc, '')
add_paragraph(doc, 'Behavioral Transformation', bold=True)
add_paragraph(doc, '')

add_paragraph(doc, 'In Every Conversation', bold=True)
add_paragraph(doc, 'Sales, product, and delivery teams now start conversations with value, not features.')
add_paragraph(doc, '')

add_paragraph(doc, 'In Project Planning', bold=True)
add_paragraph(doc, 'Every initiative must define its value story before receiving approval.')
add_paragraph(doc, '')

add_paragraph(doc, 'In Performance Reviews', bold=True)
add_paragraph(doc, 'Success is measured by customer value delivered, not just tasks completed.')
add_paragraph(doc, '')

add_paragraph(doc, 'In Celebrations', bold=True)
add_paragraph(doc, 'We celebrate validated value stories, not just product launches.')
add_paragraph(doc, '')

# Path 3: How will we get there
add_heading(doc, 'Path 3: How will we get there', 1)

# Screen 1
add_heading(doc, 'Screen 1: How we reach $100B by 2035', 2)
add_paragraph(doc, 'is achieved through systematic processes, rigorous measurement, and continuous learning.')
add_paragraph(doc, '')
add_paragraph(doc, 'Key Milestones', bold=True)
add_paragraph(doc, '')

add_paragraph(doc, 'Milestone 1 - Q1 2024', bold=True)
add_paragraph(doc, 'Launch Value Launchpad across all customer engagements. Every new project starts with defined value criteria.')
add_paragraph(doc, '')

add_paragraph(doc, 'Milestone 2 - Q3 2025', bold=True)
add_paragraph(doc, 'Achieve 50+ Level 4 value stories with validated P&L impact. Establish baseline for scaling.')
add_paragraph(doc, '')

add_paragraph(doc, 'Milestone 3 - Q4 2028', bold=True)
add_paragraph(doc, 'Deploy value tracking dashboard company-wide. Real-time visibility into all value stories across levels.')
add_paragraph(doc, '')

add_paragraph(doc, 'Milestone 4 - Q2 2032', bold=True)
add_paragraph(doc, 'Create library of 200+ replicable value blueprints. Enable rapid deployment across industries.')
add_paragraph(doc, '')

add_paragraph(doc, 'Milestone 5 - Dec 2035', bold=True)
add_paragraph(doc, 'Reach $100B in cumulative customer value. Establish Cognite as the definitive value delivery platform.')
add_paragraph(doc, '')

# Screen 2
add_heading(doc, 'Screen 2: It takes all of us', 2)
add_paragraph(doc, 'Every Cogniter has a distinct role to play in 100 × 35. Here\'s how different roles contribute to unlocking value.')
add_paragraph(doc, '')
add_paragraph(doc, 'Role Mandates', bold=True)
add_paragraph(doc, '')

add_paragraph(doc, 'Account Executive - The Architect of Value', bold=True)
add_paragraph(doc, 'Uses the Value Launchpad to frame deals around business outcomes and ensures value is baked into SOWs.')
add_paragraph(doc, '')

add_paragraph(doc, 'Product Manager - The Builder of Value', bold=True)
add_paragraph(doc, 'Prioritizes features that drive measurable outcomes and gives customers tools to track financial & value levers.')
add_paragraph(doc, '')

add_paragraph(doc, 'Project Manager - The Deliverer of Value', bold=True)
add_paragraph(doc, 'Runs projects against business goals, not just activity. Ensures milestones reflect impact on the customer\'s P&L.')
add_paragraph(doc, '')

add_paragraph(doc, 'Engineer / Data Scientist - The Designer of Value', bold=True)
add_paragraph(doc, 'Turns data and models into outcomes, framing results in terms of business KPIs and real-world decisions.')
add_paragraph(doc, '')

add_paragraph(doc, 'CBE / TAM - The Accelerator of Value', bold=True)
add_paragraph(doc, 'Monitors adoption and links technical performance back to agreed business outcomes.')
add_paragraph(doc, '')

add_paragraph(doc, 'Marketing - The Storyteller of Value', bold=True)
add_paragraph(doc, 'Converts Level 5 successes into public stories that prove Cognite\'s value assurance.')
add_paragraph(doc, '')

add_paragraph(doc, 'P&O Partner - The Enabler of Value Culture', bold=True)
add_paragraph(doc, 'Recruits and develops talent that thinks in terms of customer value.')
add_paragraph(doc, '')

add_paragraph(doc, 'Finance Partner - The Guardian of Value', bold=True)
add_paragraph(doc, 'Aligns commercial models with customer success and champions rigorous value measurement.')
add_paragraph(doc, '')

# Screen 3
add_heading(doc, 'Screen 3: It takes all of us', 2)
add_paragraph(doc, 'To support 100 × 35, Cognite is rolling out systems and processes for reporting value across the entire lifecycle – from early potential to scaled impact.')
add_paragraph(doc, '')
add_paragraph(doc, 'The Continuous Feedback Loop', bold=True)
add_paragraph(doc, '')
add_paragraph(doc, 'The Feedback Loop in Action', bold=True, italic=True)
add_paragraph(doc, '')

add_paragraph(doc, 'Launch', bold=True)
add_paragraph(doc, 'Deploy use cases to customers')
add_paragraph(doc, '')

add_paragraph(doc, 'Measure', bold=True)
add_paragraph(doc, 'Track metrics & outcomes')
add_paragraph(doc, '')

add_paragraph(doc, 'Report', bold=True)
add_paragraph(doc, 'Document value stories')
add_paragraph(doc, '')

add_paragraph(doc, 'Improve', bold=True)
add_paragraph(doc, 'Refine & optimize approach')
add_paragraph(doc, '')

# Save document
output_file = 'MoonShot_Content_Outline.docx'
doc.save(output_file)
print(f'Document created successfully: {output_file}')
